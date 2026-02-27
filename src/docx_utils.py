"""
docx_utils.py
Utility functions to open, read, write, search, and manipulate .docx files,
including image handling (extract, insert, replace).
Requires: python-docx, Pillow
"""

import os
import re
import tempfile
from copy import deepcopy
from pathlib import Path
from typing import Optional

import lxml.etree as etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image, ImageOps

EMU_PER_INCH = 914400


# ---------------------------------------------------------------------------
# Document lifecycle
# ---------------------------------------------------------------------------

def open_doc(path: str) -> Document:
    """Open an existing .docx file and return a Document object."""
    if not os.path.exists(path):
        raise FileNotFoundError(f"Document not found: {path}")
    return Document(path)


def new_doc() -> Document:
    """Create and return a blank Document."""
    return Document()


def save_doc(doc: Document, path: str) -> None:
    """Save a Document to the given path (creates parent dirs if needed)."""
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    doc.save(path)


# ---------------------------------------------------------------------------
# Reading
# ---------------------------------------------------------------------------

def get_text(doc: Document) -> str:
    """Return the full plain text of the document."""
    return "\n".join(p.text for p in doc.paragraphs)


def get_paragraphs(doc: Document) -> list:
    """Return all Paragraph objects in the document."""
    return doc.paragraphs


def get_paragraph_texts(doc: Document) -> list[str]:
    """Return a list of the text string for every paragraph."""
    return [p.text for p in doc.paragraphs]


def get_tables(doc: Document) -> list:
    """Return all Table objects in the document."""
    return doc.tables


def get_table_data(table) -> list[list[str]]:
    """Convert a Table object into a 2-D list of cell text strings."""
    return [[cell.text for cell in row.cells] for row in table.rows]


def get_styles(doc: Document) -> list[str]:
    """Return the names of all styles available in the document."""
    return [style.name for style in doc.styles]


def get_headings(doc: Document) -> list[dict]:
    """Return all heading paragraphs with their level and text."""
    headings = []
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            level = int(p.style.name.split()[-1]) if p.style.name[-1].isdigit() else 0
            headings.append({"level": level, "text": p.text, "paragraph": p})
    return headings


# ---------------------------------------------------------------------------
# Writing
# ---------------------------------------------------------------------------

def add_paragraph(doc: Document, text: str, style: Optional[str] = None) -> Paragraph:
    """Add a paragraph at the end of the document. Returns the Paragraph."""
    return doc.add_paragraph(text, style=style)


def add_heading(doc: Document, text: str, level: int = 1) -> Paragraph:
    """Add a heading at the end of the document. Returns the Paragraph."""
    return doc.add_heading(text, level=level)


def add_page_break(doc: Document) -> None:
    """Insert a page break at the end of the document."""
    doc.add_page_break()


def add_table(
    doc: Document,
    rows: int,
    cols: int,
    data: Optional[list[list[str]]] = None,
):
    """
    Add a table at the end of the document.
    If data is provided (list of lists), fill the cells with it.
    Returns the Table object.
    """
    table = doc.add_table(rows=rows, cols=cols)
    if data:
        for r_idx, row_data in enumerate(data):
            for c_idx, cell_text in enumerate(row_data):
                if r_idx < rows and c_idx < cols:
                    table.rows[r_idx].cells[c_idx].text = str(cell_text)
    return table


def set_paragraph_alignment(paragraph: Paragraph, alignment: str) -> None:
    """
    Set paragraph alignment.
    alignment: 'left' | 'center' | 'right' | 'justify'
    """
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    paragraph.alignment = mapping.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)


def set_run_format(
    run,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    font_name: Optional[str] = None,
    font_size_pt: Optional[float] = None,
    color_rgb: Optional[tuple[int, int, int]] = None,
) -> None:
    """
    Apply formatting to a Run object.
    color_rgb: (R, G, B) tuple with values 0-255.
    """
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline
    if font_name is not None:
        run.font.name = font_name
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)
    if color_rgb is not None:
        run.font.color.rgb = RGBColor(*color_rgb)


# ---------------------------------------------------------------------------
# Search & Replace
# ---------------------------------------------------------------------------

def find_in_doc(doc: Document, search_text: str, case_sensitive: bool = False) -> list:
    """Return a list of Paragraphs whose text contains search_text (regex supported)."""
    flags = 0 if case_sensitive else re.IGNORECASE
    return [p for p in doc.paragraphs if re.search(search_text, p.text, flags)]


def find_by_style(doc: Document, style_name: str) -> list:
    """Return all paragraphs that use the given style name."""
    return [p for p in doc.paragraphs if p.style.name == style_name]


def find_and_replace(
    doc: Document,
    old_text: str,
    new_text: str,
    case_sensitive: bool = False,
) -> int:
    """
    Replace all occurrences of old_text with new_text across every run.
    Note: text split across multiple runs is not matched.
    Returns the number of runs modified.
    """
    flags = 0 if case_sensitive else re.IGNORECASE
    count = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if re.search(old_text, run.text, flags):
                run.text = re.sub(old_text, new_text, run.text, flags=flags)
                count += 1
    return count


# ---------------------------------------------------------------------------
# Images — extraction
# ---------------------------------------------------------------------------

def extract_images(doc: Document, output_dir: str) -> list[str]:
    """
    Extract all images embedded in the document and save them to output_dir.
    Returns a list of saved file paths.
    """
    os.makedirs(output_dir, exist_ok=True)
    saved = []
    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.reltype:
            img_data = rel.target_part.blob
            ext = rel.target_part.content_type.split("/")[-1]
            out_path = os.path.join(output_dir, f"image_{i}.{ext}")
            with open(out_path, "wb") as f:
                f.write(img_data)
            saved.append(out_path)
    return saved


def get_inline_images(doc: Document) -> list:
    """Return all InlineShape objects (inline images) in the document."""
    return list(doc.inline_shapes)


# ---------------------------------------------------------------------------
# Images — insertion
# ---------------------------------------------------------------------------

def add_image(
    doc: Document,
    image_path: str,
    width_inches: Optional[float] = None,
    height_inches: Optional[float] = None,
):
    """
    Append an image at the end of the document.
    Returns the InlineShape object.
    """
    width = Inches(width_inches) if width_inches else None
    height = Inches(height_inches) if height_inches else None
    return doc.add_picture(image_path, width=width, height=height)


def insert_image_after_paragraph(
    doc: Document,
    paragraph: Paragraph,
    image_path: str,
    width_inches: Optional[float] = None,
    height_inches: Optional[float] = None,
) -> Paragraph:
    """
    Insert an image immediately after the given paragraph.
    Returns the new Paragraph containing the image.

    Strategy: add the picture at the end (which registers the image relationship
    correctly), then move the generated <w:p> element to the desired position.
    """
    width = Inches(width_inches) if width_inches else None
    height = Inches(height_inches) if height_inches else None

    doc.add_picture(image_path, width=width, height=height)
    pic_element = doc.paragraphs[-1]._element

    # Detach from the end and reattach right after the anchor paragraph
    pic_element.getparent().remove(pic_element)
    paragraph._element.addnext(pic_element)

    return Paragraph(pic_element, doc)


def insert_image_at_placeholder(
    doc: Document,
    placeholder_text: str,
    image_path: str,
    width_inches: Optional[float] = None,
    height_inches: Optional[float] = None,
    remove_placeholder: bool = True,
) -> bool:
    """
    Find the first paragraph containing placeholder_text and insert an image
    immediately after it. Optionally remove the placeholder paragraph afterwards.
    Returns True if the placeholder was found, False otherwise.

    Example usage:
        insert_image_at_placeholder(doc, "{{CHART}}", "chart.png", width_inches=5)
    """
    matches = find_in_doc(doc, placeholder_text)
    if not matches:
        return False
    target = matches[0]
    insert_image_after_paragraph(doc, target, image_path, width_inches, height_inches)
    if remove_placeholder:
        _delete_paragraph(target)
    return True


def replace_image_by_index(
    doc: Document,
    index: int,
    new_image_path: str,
    width_inches: Optional[float] = None,
    height_inches: Optional[float] = None,
) -> None:
    """
    Replace the inline image at position `index` (0-based) with a new image.
    Preserves the position within the document.
    """
    shapes = list(doc.inline_shapes)
    if index >= len(shapes):
        raise IndexError(f"Image index {index} out of range ({len(shapes)} images found).")

    # Navigate from InlineShape up to its enclosing w:p element
    inline_elem = shapes[index]._inline
    wp = inline_elem.getparent().getparent()  # w:r -> w:p
    para = Paragraph(wp, doc)

    insert_image_after_paragraph(doc, para, new_image_path, width_inches, height_inches)
    _delete_paragraph(para)


# ---------------------------------------------------------------------------
# Image utilities (Pillow)
# ---------------------------------------------------------------------------

def get_image_info(image_path: str) -> dict:
    """Return basic metadata of an image file using Pillow."""
    with Image.open(image_path) as img:
        return {
            "path": image_path,
            "format": img.format,
            "mode": img.mode,
            "width_px": img.width,
            "height_px": img.height,
        }


def resize_image(image_path: str, output_path: str, width_px: int, height_px: int) -> str:
    """Resize an image to the given pixel dimensions and save to output_path."""
    with Image.open(image_path) as img:
        resized = img.resize((width_px, height_px), Image.LANCZOS)
        resized.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Floating image
# ---------------------------------------------------------------------------

def _resolve_student_dir(pictures_dir, student_name: str) -> Optional[Path]:
    """Return the student's picture folder (case-insensitive match). None if not found."""
    pictures_root = Path(pictures_dir)
    student_dir = pictures_root / student_name
    if student_dir.exists():
        return student_dir
    matches = [
        d for d in pictures_root.iterdir()
        if d.is_dir() and d.name.lower() == student_name.lower()
    ]
    return matches[0] if matches else None


def find_student_picture(pictures_dir, student_name: str) -> Optional[str]:
    """Return the PORT image for the student. Returns None if not found."""
    pics = find_student_pictures(pictures_dir, student_name)
    return pics.get("PORT")


def find_student_pictures(pictures_dir, student_name: str) -> dict[str, str]:
    """
    Return a dict mapping uppercase filename stem → full path for every image
    found in the student's folder.

    Example:
        {"PORT": "/path/PORT.JPG", "INTER": "/path/INTER.JPG", ...}

    Only files whose stems exactly match one of the expected slots are included.
    Returns an empty dict if the folder or images are not found.
    """
    image_extensions = {".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tiff", ".tif"}

    student_dir = _resolve_student_dir(pictures_dir, student_name)
    if not student_dir:
        return {}

    return {
        f.stem.upper(): str(f)
        for f in student_dir.iterdir()
        if f.suffix.lower() in image_extensions
    }


def insert_floating_image(
    doc: Document,
    image_path: str,
    x_inches: float,
    y_inches: float,
    width_inches: float,
    wrap_type: str = "front",
    anchor_paragraph: Optional[Paragraph] = None,
    v_relative: str = "page",
    dist_t: float = 0.0,
    dist_b: float = 0.0,
    dist_l: float = 0.0,
    dist_r: float = 0.0,
    border_color: Optional[str] = None,
    border_pt: float = 1.0,
) -> None:
    """
    Insert an image as a floating element with absolute horizontal position and
    configurable vertical reference.
    Width is set to width_inches; height is scaled to preserve aspect ratio.
    EXIF orientation is corrected automatically.

    wrap_type — how the image interacts with text:
      "front"      — in front of text (default)
      "behind"     — behind text
      "square"     — text wraps around the bounding box
      "tight"      — text wraps tightly around the image shape
      "top_bottom" — text above and below, not beside

    anchor_paragraph — paragraph the image is embedded in.
      If None, uses the first paragraph of the document.
      The paragraph also controls where Word re-flows the image when the
      document is edited.

    v_relative — vertical reference frame for y_inches:
      "page"      — y measured from the top of the page (default)
      "paragraph" — y measured from the top of the anchor paragraph
                    (use y_inches=0 to align the image with the paragraph)

    dist_t / dist_b / dist_l / dist_r — minimum distance (in inches) between
      the image edge and surrounding text (top, bottom, left, right).
      Only meaningful for wrapping modes other than "front"/"behind".

    border_color — hex RGB string (e.g. "000000" for black). If None, no
      border is drawn.
    border_pt    — border line weight in points (only used when border_color
      is set).

    Strategy:
      1. Add the image inline at the end (registers the image relationship).
      2. Extract the <a:graphic> element (which holds the r:embed reference).
      3. Build a <wp:anchor> with the requested positioning.
      4. Inject the anchor into the target paragraph and remove the temp paragraph.
    """
    # Map wrap_type to (behindDoc value, wrap XML element)
    WRAP_MAP = {
        "front":      ("0", "<wp:wrapNone/>"),
        "behind":     ("1", "<wp:wrapNone/>"),
        "square":     ("0", "<wp:wrapSquare wrapText='bothSides'/>"),
        "tight":      ("0", "<wp:wrapTight wrapText='bothSides'/>"),
        "top_bottom": ("0", "<wp:wrapTopAndBottom/>"),
    }
    if wrap_type not in WRAP_MAP:
        raise ValueError(f"wrap_type must be one of {list(WRAP_MAP)}. Got: {wrap_type!r}")

    behind_doc, wrap_xml = WRAP_MAP[wrap_type]

    WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    A  = "http://schemas.openxmlformats.org/drawingml/2006/main"

    # Correct EXIF orientation (phone photos are often rotated in metadata only)
    with Image.open(image_path) as img:
        img = ImageOps.exif_transpose(img)
        img_w, img_h = img.size
        suffix = Path(image_path).suffix.lower() or ".jpg"
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        img.save(tmp.name)
        corrected_path = tmp.name

    width_emu  = int(width_inches * EMU_PER_INCH)
    height_emu = int(width_emu * img_h / img_w)
    x_emu      = int(x_inches * EMU_PER_INCH)
    y_emu      = int(y_inches * EMU_PER_INCH)
    dist_t_emu = int(dist_t * EMU_PER_INCH)
    dist_b_emu = int(dist_b * EMU_PER_INCH)
    dist_l_emu = int(dist_l * EMU_PER_INCH)
    dist_r_emu = int(dist_r * EMU_PER_INCH)

    # Add inline picture so python-docx registers the image relationship
    doc.add_picture(corrected_path, width=Inches(width_inches))
    last_para = doc.paragraphs[-1]
    inline = last_para._element.find(".//" + qn("wp:inline"))

    # Copy the graphic element (holds the r:embed relationship ID)
    graphic = deepcopy(inline.find(qn("a:graphic")))

    # Inject picture border into <pic:spPr> if requested
    if border_color is not None:
        PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
        PT_TO_EMU = 12700   # 1 point = 12700 EMU
        sp_pr = graphic.find(f".//{{{PIC}}}spPr")
        if sp_pr is not None:
            ln = etree.SubElement(sp_pr, f"{{{A}}}ln", w=str(int(border_pt * PT_TO_EMU)))
            solid_fill = etree.SubElement(ln, f"{{{A}}}solidFill")
            etree.SubElement(solid_fill, f"{{{A}}}srgbClr", val=border_color.lstrip("#"))

    # Pick a unique docPr id (must be unique across all drawings in the doc)
    existing_ids = {
        int(el.get("id", 0))
        for el in doc.element.body.iter()
        if el.tag == qn("wp:docPr")
    }
    doc_pr_id = max(existing_ids, default=0) + 1

    # Build the anchor element
    anchor_xml = (
        f'<wp:anchor distT="{dist_t_emu}" distB="{dist_b_emu}" distL="{dist_l_emu}" distR="{dist_r_emu}" '
        f'simplePos="0" relativeHeight="251658240" '
        f'behindDoc="{behind_doc}" '
        f'locked="0" layoutInCell="1" allowOverlap="1" '
        f'xmlns:wp="{WP}">'
        f'  <wp:simplePos x="0" y="0"/>'
        f'  <wp:positionH relativeFrom="page">'
        f'    <wp:posOffset>{x_emu}</wp:posOffset>'
        f'  </wp:positionH>'
        f'  <wp:positionV relativeFrom="{v_relative}">'
        f'    <wp:posOffset>{y_emu}</wp:posOffset>'
        f'  </wp:positionV>'
        f'  <wp:extent cx="{width_emu}" cy="{height_emu}"/>'
        f'  <wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'  {wrap_xml}'
        f'  <wp:docPr id="{doc_pr_id}" name="FloatImg{doc_pr_id}"/>'
        f'  <wp:cNvGraphicFramePr>'
        f'    <a:graphicFrameLocks xmlns:a="{A}" noChangeAspect="1"/>'
        f'  </wp:cNvGraphicFramePr>'
        f'</wp:anchor>'
    )
    anchor = etree.fromstring(anchor_xml)
    anchor.append(graphic)

    # Wrap anchor in w:r > w:drawing and inject into the target paragraph
    run = OxmlElement("w:r")
    drawing = OxmlElement("w:drawing")
    drawing.append(anchor)
    run.append(drawing)
    target_para = anchor_paragraph if anchor_paragraph is not None else doc.paragraphs[0]
    target_para._element.append(run)

    # Remove the temporary inline paragraph added by doc.add_picture()
    _delete_paragraph(last_para)

    # Clean up the orientation-corrected temp file
    os.unlink(corrected_path)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _delete_paragraph(paragraph: Paragraph) -> None:
    """Remove a paragraph element from the document body."""
    p = paragraph._element
    p.getparent().remove(p)
